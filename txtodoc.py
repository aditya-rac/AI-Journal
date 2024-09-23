import os
from docx import Document

# Define the path of the input text file and output Word document
input_file_path = r"fileName.txt"
output_doc_path = r"fileName.docx"

# Create a new Word document
doc = Document()

# Open the input text file and read its content
with open(input_file_path, 'r', encoding='utf-8') as file:
    lines = file.readlines()

# Loop through each line in the text file
for line in lines:
    # Ignore lines that contain page breaks '----'
    if '----' in line:
        continue
    
    # Check for heading levels based on the count of '#'
    if line.startswith("##"):
        heading_level = line.count("#")
        heading_text = line.strip("#").strip()
        
        # Add the heading to the Word document with the appropriate level
        doc.add_heading(heading_text, level=heading_level-1)
    else:
        # Add regular text to the document
        doc.add_paragraph(line.strip())

# Save the Word document
doc.save(output_doc_path)

print(f"Converted Word document saved to {output_doc_path}")
